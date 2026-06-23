"""
学习通试卷知识点汇总 - 将所有阶段测试的题目合并为一个 DOCX 文件
格式：题干 + 正确答案 → 知识点陈述句
例如："对某地企业进行分组，第一层按经济类型分组，第二层按企业规模分组，这样的分组属于复合分组。"
"""
import re
import os
import json
import sys
import time
from urllib.parse import urlparse, parse_qs
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ========== 阶段测试列表 ==========
EXAMS = [
    {
        "name": "阶段测试一",
        "url": "https://mooc1.xuexitong.com/exam-ans/exam/test/reVersionPaperMarkContentNew?courseId=250332561&classId=139715224&p=1&id=168547197&ut=s&newMooc=true&qbanksystem=1&qbankbackurl=%2Fexam-ans%2Fexam%2Ftest%2Flook%3FcourseId%3D250332561%26classId%3D139715224%26examId%3D9256073%26examAnswerId%3D168547197%26cpi%3D417091577&cpi=417091577&openc=649ab773508111839da4d50e1a98fe29",
    },
    {
        "name": "阶段测试二",
        "url": "https://mooc1.xuexitong.com/exam-ans/exam/test/reVersionPaperMarkContentNew?courseId=250332561&classId=139715224&p=1&id=168859383&ut=s&newMooc=true&qbanksystem=1&qbankbackurl=%2Fexam-ans%2Fexam%2Ftest%2Flook%3FcourseId%3D250332561%26classId%3D139715224%26examId%3D9388275%26examAnswerId%3D168859383%26cpi%3D417091577&cpi=417091577&openc=649ab773508111839da4d50e1a98fe29",
    },
    {
        "name": "阶段测试三",
        "url": "https://mooc1.xuexitong.com/exam-ans/exam/test/reVersionPaperMarkContentNew?courseId=250332561&classId=139715224&p=1&id=169015249&ut=s&newMooc=true&qbanksystem=1&qbankbackurl=%2Fexam-ans%2Fexam%2Ftest%2Flook%3FcourseId%3D250332561%26classId%3D139715224%26examId%3D9473677%26examAnswerId%3D169015249%26cpi%3D417091577&cpi=417091577&openc=649ab773508111839da4d50e1a98fe29",
    },
    {
        "name": "阶段测试四",
        "url": "https://mooc1.xuexitong.com/exam-ans/exam/test/reVersionPaperMarkContentNew?courseId=250332561&classId=139715224&p=1&id=169106266&ut=s&newMooc=true&qbanksystem=1&qbankbackurl=%2Fexam-ans%2Fexam%2Ftest%2Flook%3FcourseId%3D250332561%26classId%3D139715224%26examId%3D9548629%26examAnswerId%3D169106266%26cpi%3D417091577&cpi=417091577&openc=649ab773508111839da4d50e1a98fe29",
    },
    {
        "name": "阶段测试五",
        "url": "https://mooc1.xuexitong.com/exam-ans/exam/test/reVersionPaperMarkContentNew?courseId=250332561&classId=139715224&p=1&id=169265091&ut=s&newMooc=true&qbanksystem=1&qbankbackurl=%2Fexam-ans%2Fexam%2Ftest%2Flook%3FcourseId%3D250332561%26classId%3D139715224%26examId%3D9617871%26examAnswerId%3D169265091%26cpi%3D417091577&cpi=417091577&openc=649ab773508111839da4d50e1a98fe29",
    },
    {
        "name": "阶段测试六",
        "url": "https://mooc1.xuexitong.com/exam-ans/exam/test/reVersionPaperMarkContentNew?courseId=250332561&classId=139715224&p=1&id=169477023&ut=s&newMooc=true&qbanksystem=1&qbankbackurl=%2Fexam-ans%2Fexam%2Ftest%2Flook%3FcourseId%3D250332561%26classId%3D139715224%26examId%3D9747398%26examAnswerId%3D169477023%26cpi%3D417091577&cpi=417091577&openc=649ab773508111839da4d50e1a98fe29",
    },
]

OUTPUT_DOCX = "知识点汇总.docx"

# Cookie 原始数据
COOKIE_RAW = """_d	1781077358979	.xuexitong.com	/	2026-06-17T19:42:39.919Z	15						Medium
_industry	5	.xuexitong.com	/	2026-06-11T07:42:47.290Z	10						Medium
_uid	351234427	.xuexitong.com	/	2026-06-17T19:42:39.919Z	13						Medium
47728enc	2ED85A150A7295529EC5C6A4A30D5204	.xuexitong.com	/	2026-06-17T19:42:39.919Z	40						Medium
47728UID	351234427	.xuexitong.com	/	2026-06-17T19:42:39.919Z	17						Medium
47728userinfo	47cfb0decaadf37d2d5350157aca49f2c49d67c0c30ca5047c5a963e85f1109990df822afcd77ec5c98bf515e91478716f0a630ffff96640e7fafd565af53bf2	.xuexitong.com	/	2026-06-17T19:42:39.919Z	141						Medium
cx_p_token	c9048e20ecc932e0034fbd412f11b86c	.xuexitong.com	/	2026-06-17T19:42:39.920Z	42						Medium
DSSTASH_LOG	C_38-UN_341-US_351234427-T_1781077358981	.xuexitong.com	/	2026-06-17T19:42:39.920Z	51						Medium
fid	47728	.xuexitong.com	/	2026-07-10T07:42:39.919Z	8						Medium
jrose	72BAFE46D6F46F3B6762AF2B5D2CF22A.mooc-2010510164-qnlr9	mooc1.xuexitong.com	/	Session	59	✓	✓				Medium
jrose	DA12BE241C22F9E01422A839886624BF.mooc-exam-1805053833-95hbf	mooc1.xuexitong.com	/exam-ans	Session	64	✓	✓				Medium
jrosehead	AAA23DB76E1EB6D06DF80CAFD380C16C.mooc-portal-2961312520-9l52z	mooc1.xuexitong.com	/	Session	70	✓	✓				Medium
k8s	1781078134.777.20562.258644	stat2-ans.xuexitong.com	/	Session	30	✓					Medium
k8s	1781077405.317.1333.952044	mooc1.xuexitong.com	/	Session	29	✓					Medium
k8sexam	1781077488.649.17415.258365	mooc1.xuexitong.com	/exam-ans	Session	34	✓					Medium
p_auth_token	eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1aWQiOiIzNTEyMzQ0MjciLCJsb2dpblRpbWUiOjE3ODEwNzczNTg5ODEsImV4cCI6MTc4MTY4MjE1OH0.QZ0_NUexFkNPcNOa-q5mLHXGQA2u_2PAqJN4d7GROTY	.xuexitong.com	/	2026-06-17T19:42:39.920Z	176	✓					Medium
route	c010ccedb771f8b7c7793c67ee1d2aae	mooc1.xuexitong.com	/	Session	37						Medium
route	bca6486eee9aca907e6257b7921729c3	stat2-ans.xuexitong.com	/	Session	37						Medium
source	""	.xuexitong.com	/	Session	8						Medium
spaceFid	47728	.xuexitong.com	/	Session	13						Medium
spaceRoleId	""	.xuexitong.com	/	Session	13						Medium
thirdRegist	0	.xuexitong.com	/	Session	12						Medium
uf	d9387224d3a6095b452718da10d186f2bc9e8f9c201149fe5e3f0041e2b172e1b13c4f654d35f4d5001ed2a0ae69c9bec7ea6fb664318d21c49d67c0c30ca5043ad701c8b4cc548c0234d89f51c3dccfe34d59cb42313048713028f1ec42bf71b1188854805578cc9077eb294bc1b943e09bccdbc39a98e362112372be26ec5920690c991987d88df9e630debd70cfb54df7ff280fcb29d10d8a4c92b12beb4b6a4c161ae4de1d619f75b8dbc17d86506250480410be0c44e7fafd565af53bf2	.xuexitong.com	/	2026-06-17T19:42:39.919Z	386						Medium
UID	351234427	.xuexitong.com	/	2026-06-17T19:42:39.919Z	12						Medium
vc3	GTkh76KeSplbltaoGEsp1jTxr6bv0Z85DPHZ4qafXqNvoVm34B6z2ndSWDmwf8UiokoC5LhWNQAZwzVtMNmjmYZvmAAbjr4q4PHSN7NGfZ7Y3ybufgCP43S7hUA9SQsEK%2B%2BE%2F%2Bme54JPc%2BgnUNfd7W35yn01rUWWD1ioXGvFOrw%3D3ad975d71c719d2e5512b20df897ac37	.xuexitong.com	/	2026-06-17T19:42:39.919Z	219	✓					Medium
xxtenc	679d2ad6b066649e74a9bdad8728ecad	.xuexitong.com	/	2026-07-10T07:42:39.920Z	38						Medium"""

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9",
}


def parse_cookie_jar(raw: str):
    jar = requests.cookies.RequestsCookieJar()
    for line in raw.strip().split("\n"):
        parts = line.split("\t")
        if len(parts) >= 4:
            name = parts[0].strip()
            value = parts[1].strip().strip('"')
            if value == '""':
                value = ""
            jar.set(name, value, domain=parts[2].strip(), path=parts[3].strip())
    return jar


def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def download_image(session: requests.Session, url: str, save_dir: str, prefix: str, idx: int) -> str:
    """下载图片，带重试机制（最多3次），单次超时15秒"""
    parsed = urlparse(url)
    orig_name = os.path.basename(parsed.path)
    if not orig_name or '.' not in orig_name:
        ext = '.png'
    else:
        ext = os.path.splitext(orig_name)[1]
        if len(ext) > 5:
            ext = '.png'
    filename = f"{prefix}_图{idx}{ext}"
    filepath = os.path.join(save_dir, filename)

    # 如果已下载过，跳过
    if os.path.exists(filepath):
        return filename

    max_retries = 3
    for attempt in range(1, max_retries + 1):
        try:
            resp = session.get(url, timeout=(10, 15))  # (connect_timeout, read_timeout)
            resp.raise_for_status()
            with open(filepath, "wb") as f:
                f.write(resp.content)
            return filename
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            if attempt < max_retries:
                wait = attempt * 2
                print(f"    [!] 图片下载重试 {attempt}/{max_retries}（{wait}s后）: {url[:60]}...")
                time.sleep(wait)
            else:
                print(f"    [!] 图片下载失败（已重试{max_retries}次）: {url[:60]}... - {e}")
                return ""
        except requests.exceptions.HTTPError as e:
            print(f"    [!] 图片下载HTTP错误: {url[:60]}... - {e}")
            return ""
        except KeyboardInterrupt:
            raise  # 向上传递，让外层处理
        except Exception as e:
            print(f"    [!] 图片下载失败: {url[:60]}... - {e}")
            return ""
    return ""


def parse_question_li(li) -> dict:
    result = {}
    full_text = li.get_text(separator="\n", strip=True)
    result["raw_text"] = full_text

    # 图片
    images = []
    for img in li.select("img"):
        src = img.get("src", "")
        if src and ("ananas.chaoxing.com" in src or "p.ananas" in src):
            images.append(src)
    if images:
        result["images"] = images

    # 正确答案（CSS选择器: 选择题为 span.rightAnswerContent，填空题为 dd.rightAnswerContent）
    right_els = li.select("span.rightAnswerContent, dd.rightAnswerContent")
    if right_els:
        answers = [el.get_text(strip=True) for el in right_els]
        result["correct_answer"] = "；".join(answers)

    # 我的答案（兜底用）
    stu_els = li.select("span.stuAnswerContent, .stuAnswerContent")
    if stu_els:
        my_answers = [el.get_text(strip=True) for el in stu_els]
        result["my_answer"] = "；".join(my_answers)

    # 文本解析
    first_line = full_text.split("\n")[0] if full_text else ""
    m = re.match(r'^(\d+)\.', first_line)
    if m:
        result["number"] = int(m.group(1))

    m = re.search(r'[（(]([^）)]+)[）)]', full_text)
    if m:
        result["type"] = m.group(1)

    # 提取题干 — 题型+分值后面的内容，直到第一个选项前
    m = re.search(r'[\d]+\.\s*[（(][^）)]+[）)]\s*(.+)', full_text)
    if m:
        stem = m.group(1)
        stem = re.split(r'\s+[A-F][.．、]\s*', stem)[0]
        result["stem"] = clean_text(stem)

    # 提取选项（修复: 用 split 方式确保 D 选项不漏掉）
    options = {}
    # 找到"A."之后的所有文本
    opt_start = re.search(r'\s+A[.．、]\s*', full_text)
    if opt_start:
        opt_text = full_text[opt_start.start():]
        # 用正则切分 A.xxx B.xxx C.xxx D.xxx ...
        parts = re.split(r'\s+([A-F])[.．、]\s*', opt_text)
        # parts[0] 为空, parts[1]='A', parts[2]='xxx', parts[3]='B', parts[4]='xxx', ...
        i = 1
        while i < len(parts) - 1:
            letter = parts[i]
            text = parts[i + 1]
            i += 2
            # 截断到 "我的答案" 或 "正确答案" 之前
            text = re.split(r'我的答案|正确答案', text)[0]
            text = clean_text(text)
            if text:
                options[letter] = text

    if options:
        result["options"] = options

    # 正则兜底：正确答案
    if not result.get("correct_answer"):
        m = re.search(r'正确答案\s*[:：]\s*(.+?)(?:\n|$)', full_text)
        if m:
            result["correct_answer"] = m.group(1).strip()

    return result


def extract_from_page(html: str) -> list[dict]:
    soup = BeautifulSoup(html, "html.parser")
    items = soup.select(".questionLi")
    if not items:
        items = soup.select("div[class*='question']")
    questions = []
    for i, item in enumerate(items, 1):
        try:
            q = parse_question_li(item)
            q["index"] = i
            questions.append(q)
        except Exception as e:
            print(f"    [!] 解析第 {i} 题失败: {e}")
    return questions


def make_knowledge_point(q: dict) -> str:
    """
    将题目+正确答案组合为知识点陈述句。

    单选题/多选题/判断题：用选项文字替换答案字母
      例："题干...属于( )。A.简单分组 B.复合分组" + 答案"B" → "题干...属于复合分组。"
    填空题：把答案填入题干的 ____ 处
      例："相关系数分为____相关、____相关。" + 答案"(1)完全 (2)不完全" → "相关系数分为完全相关、不完全相关。"
    简答题：题干后追加 "参考答案: 答案内容"
    """
    stem = q.get("stem", "")
    ans = q.get("correct_answer", "")
    my_ans = q.get("my_answer", "")
    qtype = q.get("type", "")
    options = q.get("options", {})

    if not stem:
        stem = q.get("raw_text", "")[:100]

    # 没有正确答案时，用我的答案
    if not ans:
        ans = my_ans

    if not ans:
        # 没有正确答案，只返回题干
        return f"{stem}（暂无答案）"

    # 判断题型
    is_choice = any(t in qtype for t in ["单选", "多选"])
    is_judge = "判断" in qtype
    is_fill = "填空" in qtype

    if is_choice and options:
        # 选择题：把答案字母替换为选项内容
        # 答案可能是 "B" 或 "BD" 或 "B；D"
        ans_letters = re.findall(r'[A-F]+', ans)
        ans_texts = []
        for letter_str in ans_letters:
            for ch in letter_str:
                if ch in options:
                    ans_texts.append(options[ch])

        if ans_texts:
            answer_text = "、".join(ans_texts)
            # 去掉题干末尾的 ( ) 或 (  ) 等填空括号
            stem = re.sub(r'[（(]\s*[）)]\s*$', '', stem)
            # 去掉末尾的标点后拼接答案
            stem = stem.rstrip('。，,；; ')
            return f"{stem}{answer_text}。"

    if is_judge and options:
        # 判断题：根据答案选对应选项文字
        ans_letter = ans.strip().upper()
        if ans_letter in options:
            return f"{stem}{options[ans_letter]}。"
        # 如果答案就是选项文字
        for letter, text in options.items():
            if ans in text or text in ans:
                return f"{stem}{text}。"

    if is_fill:
        # 填空题：把答案填入 ____ 处
        # 答案格式: "(1) 完全 (2) 不完全" 或 "完全；不完全"
        ans_parts = re.split(r'[；;]', ans)
        # 去掉编号如 "(1)" "(2)"
        clean_answers = [re.sub(r'^\(\d+\)\s*', '', a).strip() for a in ans_parts]

        result = stem
        for ca in clean_answers:
            # 替换第一个 ____
            result = re.sub(r'_{2,}', ca, result, count=1)

        if result == stem:
            # 如果没有 ____ 占位符，直接追加
            result = f"{stem}{'；'.join(clean_answers)}。"
        return result

    # 简答题和其他
    return f"{stem}\n参考答案: {ans}"


def main():
    print("=" * 60)
    print("  学习通知识点汇总生成器")
    print(f"  共 {len(EXAMS)} 个阶段测试")
    print("=" * 60)

    session = requests.Session()
    session.cookies.update(parse_cookie_jar(COOKIE_RAW))
    session.headers.update(HEADERS)
    # 配置重试策略：遇到超时/连接错误自动重试
    retry_strategy = Retry(
        total=2,
        backoff_factor=1,
        status_forcelist=[500, 502, 503, 504],
        allowed_methods=["GET"],
    )
    adapter = HTTPAdapter(max_retries=retry_strategy, pool_connections=5, pool_maxsize=10)
    session.mount("https://", adapter)
    session.mount("http://", adapter)

    base_dir = "."

    # 汇总所有题目
    all_questions = []  # [(exam_name, question_dict), ...]
    interrupted = False

    for exam in EXAMS:
        if interrupted:
            break
        name = exam["name"]
        url = exam["url"]
        image_dir = os.path.join(base_dir, name, f"{name}_图片")

        print(f"\n  [{name}] 正在爬取...")
        try:
            params = parse_qs(urlparse(url).query)
            exam_id = params.get("examId", [""])[0]
            exam_answer_id = params.get("id", [""])[0]
            referer = f"https://mooc1.xuexitong.com/exam-ans/exam/test/look?courseId=250332561&classId=139715224&examId={exam_id}&examAnswerId={exam_answer_id}&cpi=417091577"
            session.headers["Referer"] = referer

            resp = session.get(url, timeout=30)
            resp.encoding = "utf-8"
            if resp.status_code != 200:
                print(f"    [!] HTTP {resp.status_code}，跳过")
                continue

            questions = extract_from_page(resp.text)
            print(f"    [+] 提取 {len(questions)} 道题目")

            # 下载图片
            img_count = sum(len(q.get("images", [])) for q in questions)
            if img_count > 0:
                os.makedirs(image_dir, exist_ok=True)
                downloaded = 0
                skipped_by_user = False
                for qi, q in enumerate(questions):
                    if skipped_by_user:
                        break
                    imgs = q.get("images", [])
                    if not imgs:
                        continue
                    local_names = []
                    for i, img_url in enumerate(imgs, 1):
                        if skipped_by_user:
                            break
                        try:
                            fname = download_image(session, img_url, image_dir, f"{name}_第{q['index']}题", i)
                        except KeyboardInterrupt:
                            print(f"\n    [!] 用户中断，跳过剩余图片")
                            skipped_by_user = True
                            break
                        if fname:
                            local_names.append(fname)
                            downloaded += 1
                            print(f"    [+] 图片 {downloaded}/{img_count}: {fname}")
                    q["local_images"] = local_names
                    q["image_dir"] = image_dir
                print(f"    [+] 图片下载完成: {downloaded}/{img_count} 张")

            for q in questions:
                q["exam_name"] = name
                all_questions.append(q)

        except KeyboardInterrupt:
            print(f"\n    [!] 用户中断，跳过 {name} 剩余部分")
            interrupted = True
            # 保留已处理的题目
            for q in questions:
                q["exam_name"] = name
                all_questions.append(q)
            break
        except Exception as e:
            print(f"    [!] {name} 处理出错: {e}")
            continue

    # ========== 生成汇总 DOCX ==========
    if not all_questions:
        print("\n[!] 没有提取到任何题目，退出")
        return

    print(f"\n{'=' * 60}")
    print(f"  生成知识点汇总文档，共 {len(all_questions)} 题")
    print(f"{'=' * 60}")

    doc = Document()

    # 默认字体 + 紧凑段落间距
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    # 标题
    title = doc.add_heading('学习通知识点汇总', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'共 {len(all_questions)} 个知识点（覆盖 {len(EXAMS)} 个阶段测试）').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # 按阶段测试分组输出，紧凑排版，无多余空行
    current_exam = None
    q_counter = 0
    for q in all_questions:
        exam_name = q.get("exam_name", "")

        # 新阶段测试标题
        if exam_name != current_exam:
            current_exam = exam_name
            doc.add_heading(f'{exam_name}', level=2)

        q_counter += 1

        # 知识点语句（不输出题型标签）
        kp = make_knowledge_point(q)

        # 知识点正文: 编号 + 内容，紧凑排列
        p_body = doc.add_paragraph()
        run_num = p_body.add_run(f"{q_counter}. ")
        run_num.bold = True
        run_num.font.size = Pt(11)
        run_body = p_body.add_run(kp)
        run_body.font.size = Pt(11)

        # 图片（如果有的话，紧跟知识点后）
        local_imgs = q.get("local_images", [])
        img_dir = q.get("image_dir", "")
        for img_name in local_imgs[:1]:  # 最多嵌1张图，避免文档过大
            img_path = os.path.join(img_dir, img_name)
            if os.path.exists(img_path):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Inches(4.0))
                except Exception:
                    pass

    # 设置窄页边距，更紧凑
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    doc.save(OUTPUT_DOCX)
    print(f"\n[+] 汇总文档已保存: {OUTPUT_DOCX}")
    print(f"[+] 共 {q_counter} 个知识点")
    print("\n" + "=" * 60)
    print("  生成完成!")
    print("=" * 60)


if __name__ == "__main__":
    main()
